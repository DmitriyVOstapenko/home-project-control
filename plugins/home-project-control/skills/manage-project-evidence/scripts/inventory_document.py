#!/usr/bin/env python3
"""Inventory the real coverage units of one indexed document before full reading."""

from __future__ import annotations

import argparse
import importlib.util
import json
import sys
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from inspect_project import is_linklike, require_ready_project


INVENTORY_NAMESPACE = uuid.UUID("5123bcb2-5e82-4a04-8751-52a9f44f768a")
METHOD_VERSION = "1.0"


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def module_available(name: str) -> bool:
    return importlib.util.find_spec(name) is not None


def pdf_inventory(path: Path) -> tuple[str, list[object], dict, list[str], list[str]]:
    if not module_available("pypdf"):
        return "blocked", [], {}, ["Python module pypdf is unavailable"], ["Use a PDF-capable tool"]
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        result = reader.decrypt("")
        if not result:
            return "blocked", [], {"encrypted": True}, ["PDF requires a password"], []
    pages = []
    reading_requirements: list[str] = []
    for number, page in enumerate(reader.pages, 1):
        try:
            text_chars = len(page.extract_text() or "")
        except Exception:
            text_chars = 0
        try:
            image_count = len(page.images)
        except Exception:
            image_count = None
        if text_chars == 0:
            reading_requirements.append(f"page {number}: OCR or visual reading required")
        if image_count:
            reading_requirements.append(f"page {number}: visual verification required")
        pages.append({"page": number, "extractable_text_chars": text_chars, "image_count": image_count})
    return "complete", list(range(1, len(pages) + 1)), {"pages": pages}, [], reading_requirements


def workbook_inventory(path: Path) -> tuple[str, list[object], dict, list[str], list[str]]:
    if path.suffix.lower() == ".xls":
        return "blocked", [], {}, ["Legacy .xls inventory is not supported"], ["Convert a copy to .xlsx without replacing the source"]
    if not module_available("openpyxl"):
        return "blocked", [], {}, ["Python module openpyxl is unavailable"], ["Use a spreadsheet-capable tool"]
    from openpyxl import load_workbook

    keep_vba = path.suffix.lower() == ".xlsm"
    workbook = load_workbook(path, read_only=False, data_only=False, keep_vba=keep_vba)
    sheets = []
    requirements: list[str] = []
    for sheet in workbook.worksheets:
        formula_count = 0
        comment_count = 0
        for row in sheet.iter_rows():
            for cell in row:
                if cell.data_type == "f":
                    formula_count += 1
                if cell.comment is not None:
                    comment_count += 1
        hidden_rows = [index for index, dimension in sheet.row_dimensions.items() if dimension.hidden]
        hidden_columns = [index for index, dimension in sheet.column_dimensions.items() if dimension.hidden]
        if sheet.sheet_state != "visible":
            requirements.append(f"sheet {sheet.title}: inspect hidden sheet")
        if hidden_rows or hidden_columns:
            requirements.append(f"sheet {sheet.title}: inspect hidden rows or columns")
        if formula_count:
            requirements.append(f"sheet {sheet.title}: inspect formulas and displayed values")
        sheets.append(
            {
                "name": sheet.title,
                "state": sheet.sheet_state,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "merged_ranges": [str(value) for value in sheet.merged_cells.ranges],
                "hidden_rows": hidden_rows,
                "hidden_columns": hidden_columns,
                "formula_count": formula_count,
                "comment_count": comment_count,
            }
        )
    defined_names = sorted(str(name) for name in workbook.defined_names)
    features = {
        "sheets": sheets,
        "defined_names": defined_names,
        "external_link_count": len(getattr(workbook, "_external_links", [])),
        "keep_vba": keep_vba,
    }
    workbook.close()
    return "complete", [sheet["name"] for sheet in sheets], features, [], requirements


def docx_inventory(path: Path) -> tuple[str, list[object], dict, list[str], list[str]]:
    if not module_available("docx"):
        return "blocked", [], {}, ["Python module python-docx is unavailable"], ["Use a document-capable tool"]
    from docx import Document

    document = Document(path)
    with zipfile.ZipFile(path) as archive:
        names = sorted(archive.namelist())
    relevant_parts = [
        name
        for name in names
        if name.startswith("word/")
        and (
            name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
            or name in {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
        )
    ]
    requirements = [f"inspect OOXML part {name}" for name in relevant_parts if name != "word/document.xml"]
    features = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "relevant_parts": relevant_parts,
        "embedded_object_parts": [name for name in names if name.startswith("word/embeddings/")],
    }
    units = relevant_parts or ["word/document.xml"]
    return "complete", units, features, [], requirements


def image_inventory(path: Path) -> tuple[str, list[object], dict, list[str], list[str]]:
    if not module_available("PIL"):
        return "blocked", [], {}, ["Python module Pillow is unavailable"], ["Use an image-capable tool"]
    from PIL import Image

    with Image.open(path) as image:
        frame_count = getattr(image, "n_frames", 1)
        features = {
            "format": image.format,
            "width": image.width,
            "height": image.height,
            "mode": image.mode,
            "frame_count": frame_count,
        }
    return "complete", list(range(1, frame_count + 1)), features, [], ["visual reading and OCR verification required"]


def text_inventory(path: Path) -> tuple[str, list[object], dict, list[str], list[str]]:
    try:
        lines = path.read_text(encoding="utf-8").splitlines()
    except UnicodeError:
        return "blocked", [], {}, ["Text is not UTF-8"], ["Determine the source encoding without changing the file"]
    units: list[object] = list(range(1, len(lines) + 1)) or ["empty-document"]
    return "complete", units, {"line_count": len(lines)}, [], []


def inspect_file(path: Path) -> tuple[str, list[object], dict, list[str], list[str], str]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        values = pdf_inventory(path)
        method = "pypdf"
    elif extension in {".xlsx", ".xlsm", ".xltx", ".xltm", ".xls"}:
        values = workbook_inventory(path)
        method = "openpyxl"
    elif extension == ".docx":
        values = docx_inventory(path)
        method = "python-docx+zipfile"
    elif extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        values = image_inventory(path)
        method = "Pillow"
    elif extension in {".txt", ".md", ".csv", ".tsv", ".json", ".xml"}:
        values = text_inventory(path)
        method = "utf8-lines"
    else:
        values = ("blocked", [], {}, [f"Unsupported document format: {extension or '(none)'}"], [], "none")
        return values
    return (*values, method)


def active_document(root: Path, relative_path: str) -> tuple[dict, dict]:
    normalized = Path(relative_path).as_posix()
    registry = json.loads((root / ".home-control" / "documents.json").read_text(encoding="utf-8"))
    for document in registry.get("items", []):
        if isinstance(document, dict) and document.get("relative_path") == normalized and document.get("status") == "active":
            versions = document.get("versions", [])
            if not isinstance(versions, list) or not versions:
                raise ValueError(f"Indexed document has no version history: {normalized}")
            current = max(versions, key=lambda item: item.get("version", 0))
            return document, current
    raise ValueError(f"No active indexed document at relative path: {normalized}")


def append_inventory_atomic(path: Path, record: dict) -> bool:
    existing_text = path.read_text(encoding="utf-8")
    for raw in existing_text.splitlines():
        if not raw.strip():
            continue
        value = json.loads(raw)
        if not isinstance(value, dict):
            raise ValueError("document_inventories.jsonl contains a non-object record")
        if value.get("inventory_id") == record["inventory_id"]:
            stable_existing = {key: item for key, item in value.items() if key != "recorded_at_utc"}
            stable_record = {key: item for key, item in record.items() if key != "recorded_at_utc"}
            if stable_existing != stable_record:
                raise ValueError("An inventory with the same deterministic ID has different content")
            return False
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temporary.write_text(existing_text + json.dumps(record, ensure_ascii=False) + "\n", encoding="utf-8")
        temporary.replace(path)
    finally:
        if temporary.exists():
            temporary.unlink()
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("project_dir", type=Path)
    parser.add_argument("relative_path", help="Indexed path relative to the project root")
    parser.add_argument("--apply", action="store_true", help="Append the inventory after preview")
    args = parser.parse_args()

    root = require_ready_project(args.project_dir)
    document, version = active_document(root, args.relative_path)
    source = root / str(document["relative_path"])
    if is_linklike(source) or not source.is_file() or root not in source.resolve().parents:
        raise ValueError("Indexed source is missing, linked, or escapes the project")
    status, expected_units, features, blockers, requirements, method = inspect_file(source)
    record = {
        "source_document_id": document["document_id"],
        "document_version": version["version"],
        "sha256": version["sha256"],
        "relative_path": document["relative_path"],
        "format": source.suffix.lower(),
        "status": status,
        "expected_units": expected_units,
        "features": features,
        "blockers": blockers,
        "reading_requirements": requirements,
        "method": method,
        "method_version": METHOD_VERSION,
    }
    signature = json.dumps(record, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    record = {
        "inventory_id": "DI-" + uuid.uuid5(INVENTORY_NAMESPACE, signature).hex[:20],
        **record,
        "recorded_at_utc": utc_now(),
    }
    result = {"mode": "preview", "would_append": True, "inventory": record}
    if args.apply:
        appended = append_inventory_atomic(root / ".home-control" / "document_inventories.jsonl", record)
        result.update({"mode": "applied", "appended": appended, "would_append": False})
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (OSError, UnicodeError, ValueError, json.JSONDecodeError, zipfile.BadZipFile) as exc:
        print(f"Document inventory failed: {exc}", file=sys.stderr)
        raise SystemExit(2)
